"""No-retention document sanitation and native extraction.

Raw uploads are accepted only long enough to inspect and extract in memory. The
returned object deliberately contains no original bytes, path, or EXIF data.
"""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass
from typing import Any

import fitz
from docx import Document
from openpyxl import load_workbook
from PIL import Image, UnidentifiedImageError

from app.services.upload_policy import (
    MAX_CSV_CELL_CHARS,
    MAX_CSV_COLUMNS,
    MAX_CSV_ROWS,
    MAX_EXTRACTED_TEXT_CHARS,
    MAX_IMAGE_DIMENSION,
    MAX_IMAGE_PIXELS,
    MAX_PDF_IMAGES,
    MAX_PDF_PAGES,
    MAX_PDF_XREFS,
    MAX_WORKBOOK_CELLS,
    MAX_WORKBOOK_SHEETS,
    DocumentSafetyError,
    SanitizedUpload,
    sanitize_upload,
)

Image.MAX_IMAGE_PIXELS = MAX_IMAGE_PIXELS


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    mime_type: str
    size_bytes: int
    status: str
    text: str
    normalized_fields: dict[str, Any]
    findings: list[dict[str, Any]]
    confidence: float
    error_code: str | None = None
    error_message: str | None = None
    parser_provenance: str = "native"


def _bounded_text(parts: list[str]) -> str:
    joined = "\n".join(part for part in parts if part).strip()
    return joined[:MAX_EXTRACTED_TEXT_CHARS]


def _safe_cell(value: Any) -> str:
    text = "" if value is None else str(value)
    # Keep formula-like values inert if displayed or later exported by the UI.
    if text.startswith(("=", "+", "-", "@")):
        text = "'" + text
    return text[:MAX_CSV_CELL_CHARS]


def _extract_pdf(upload: SanitizedUpload) -> tuple[str, str | None, str | None]:
    try:
        document = fitz.open(stream=upload.data, filetype="pdf")
    except Exception as exc:
        raise DocumentSafetyError("EXTRACTION_FAILED", "The PDF could not be parsed safely.") from exc
    try:
        if document.is_encrypted:
            raise DocumentSafetyError("ENCRYPTED_FILE", "Password-protected PDFs are not accepted.")
        if document.page_count > MAX_PDF_PAGES:
            raise DocumentSafetyError("CONTENT_LIMIT_EXCEEDED", "PDFs may contain at most 150 pages.")
        if document.embfile_count() > 0:
            raise DocumentSafetyError("UNSAFE_PDF", "PDFs with embedded files are not accepted.")
        if document.xref_length() > MAX_PDF_XREFS:
            raise DocumentSafetyError("CONTENT_LIMIT_EXCEEDED", "The PDF contains too many internal objects.")
        parts: list[str] = []
        image_count = 0
        for page in document:
            image_count += len(page.get_images(full=True))
            if image_count > MAX_PDF_IMAGES:
                raise DocumentSafetyError("CONTENT_LIMIT_EXCEEDED", "The PDF contains too many images.")
            parts.append(page.get_text("text", sort=True))
            if sum(len(part) for part in parts) >= MAX_EXTRACTED_TEXT_CHARS:
                break
        text = _bounded_text(parts)
        if not text:
            return text, "OCR_UNAVAILABLE", "This scanned PDF needs an OCR worker, which is not enabled in this deployment."
        return text, None, None
    finally:
        document.close()


def _extract_docx(upload: SanitizedUpload) -> str:
    try:
        document = Document(io.BytesIO(upload.data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(_safe_cell(cell.text) for cell in row.cells))
                if sum(len(part) for part in parts) >= MAX_EXTRACTED_TEXT_CHARS:
                    return _bounded_text(parts)
        return _bounded_text(parts)
    except DocumentSafetyError:
        raise
    except Exception as exc:
        raise DocumentSafetyError("EXTRACTION_FAILED", "The DOCX file could not be parsed safely.") from exc


def _extract_xlsx(upload: SanitizedUpload) -> str:
    try:
        workbook = load_workbook(io.BytesIO(upload.data), read_only=True, data_only=True, keep_links=False)
        if len(workbook.worksheets) > MAX_WORKBOOK_SHEETS:
            raise DocumentSafetyError("CONTENT_LIMIT_EXCEEDED", "The workbook contains too many sheets.")
        parts: list[str] = []
        cells_seen = 0
        for worksheet in workbook.worksheets:
            parts.append(f"[{worksheet.title}]")
            for row in worksheet.iter_rows(values_only=True):
                cells_seen += len(row)
                if cells_seen > MAX_WORKBOOK_CELLS:
                    raise DocumentSafetyError("CONTENT_LIMIT_EXCEEDED", "The workbook contains too many cells.")
                parts.append(" | ".join(_safe_cell(value) for value in row))
                if sum(len(part) for part in parts) >= MAX_EXTRACTED_TEXT_CHARS:
                    return _bounded_text(parts)
        return _bounded_text(parts)
    except DocumentSafetyError:
        raise
    except Exception as exc:
        raise DocumentSafetyError("EXTRACTION_FAILED", "The XLSX file could not be parsed safely.") from exc


def _extract_csv(upload: SanitizedUpload) -> str:
    try:
        text = upload.data.decode("utf-8-sig", errors="strict")
    except UnicodeDecodeError as exc:
        raise DocumentSafetyError("UNSUPPORTED_ENCODING", "CSV files must use UTF-8 encoding.") from exc
    previous_limit = csv.field_size_limit()
    csv.field_size_limit(MAX_CSV_CELL_CHARS)
    try:
        reader = csv.reader(io.StringIO(text, newline=""))
        parts: list[str] = []
        for index, row in enumerate(reader, start=1):
            if index > MAX_CSV_ROWS:
                raise DocumentSafetyError("CONTENT_LIMIT_EXCEEDED", "CSV files may contain at most 100,000 rows.")
            if len(row) > MAX_CSV_COLUMNS:
                raise DocumentSafetyError("CONTENT_LIMIT_EXCEEDED", "CSV files may contain at most 50 columns.")
            parts.append(" | ".join(_safe_cell(value) for value in row))
            if sum(len(part) for part in parts) >= MAX_EXTRACTED_TEXT_CHARS:
                return _bounded_text(parts)
        return _bounded_text(parts)
    except csv.Error as exc:
        raise DocumentSafetyError("EXTRACTION_FAILED", "The CSV file could not be parsed safely.") from exc
    finally:
        csv.field_size_limit(previous_limit)


def _extract_image(upload: SanitizedUpload) -> tuple[str, str, str]:
    try:
        with Image.open(io.BytesIO(upload.data)) as image:
            image.verify()
        with Image.open(io.BytesIO(upload.data)) as image:
            if image.format not in {"PNG", "JPEG"}:
                raise DocumentSafetyError("TYPE_MISMATCH", "The image format does not match the document type.")
            width, height = image.size
            if width > MAX_IMAGE_DIMENSION or height > MAX_IMAGE_DIMENSION or width * height > MAX_IMAGE_PIXELS:
                raise DocumentSafetyError("CONTENT_LIMIT_EXCEEDED", "Images may not exceed 25 megapixels or 8,000 pixels per side.")
            if getattr(image, "n_frames", 1) != 1:
                raise DocumentSafetyError("UNSUPPORTED_TYPE", "Animated images are not accepted.")
        return "", "OCR_UNAVAILABLE", "Image OCR is not enabled in this deployment."
    except DocumentSafetyError:
        raise
    except (UnidentifiedImageError, OSError, Image.DecompressionBombError) as exc:
        raise DocumentSafetyError("EXTRACTION_FAILED", "The image could not be decoded safely.") from exc


_LABEL_PATTERNS = {
    "invoice_number": re.compile(r"(?:invoice\s*(?:number|no\.?))\s*[:#-]?\s*([A-Z0-9-]{3,64})", re.IGNORECASE),
    "purchase_order": re.compile(r"(?:purchase\s*order|po\s*(?:number|no\.?))\s*[:#-]?\s*([A-Z0-9-]{3,64})", re.IGNORECASE),
}


def _structured_suggestions(text: str) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    fields: dict[str, Any] = {}
    findings: list[dict[str, Any]] = []
    for field, pattern in _LABEL_PATTERNS.items():
        match = pattern.search(text)
        if match:
            value = match.group(1)
            fields[field] = value
            findings.append({"field": field, "value": value, "confidence": 0.7, "provenance": "native_text_label"})
    return fields, findings


def sanitize_and_extract(filename: str | None, mime_type: str | None, data: bytes) -> ExtractedDocument:
    upload = sanitize_upload(filename, mime_type, data)
    error_code: str | None = None
    error_message: str | None = None
    provenance = "native"
    if upload.mime_type == "application/pdf":
        text, error_code, error_message = _extract_pdf(upload)
    elif upload.mime_type.endswith("wordprocessingml.document"):
        text = _extract_docx(upload)
    elif upload.mime_type.endswith("spreadsheetml.sheet"):
        text = _extract_xlsx(upload)
    elif upload.mime_type == "text/csv":
        text = _extract_csv(upload)
    else:
        text, error_code, error_message = _extract_image(upload)
        provenance = "image_metadata_only"
    fields, findings = _structured_suggestions(text)
    if error_code == "OCR_UNAVAILABLE":
        return ExtractedDocument(
            filename=upload.filename,
            mime_type=upload.mime_type,
            size_bytes=len(upload.data),
            status="needs_ocr",
            text=text,
            normalized_fields=fields,
            findings=findings,
            confidence=0.0,
            error_code=error_code,
            error_message=error_message,
            parser_provenance=provenance,
        )
    return ExtractedDocument(
        filename=upload.filename,
        mime_type=upload.mime_type,
        size_bytes=len(upload.data),
        status="extracted",
        text=text,
        normalized_fields=fields,
        findings=findings,
        confidence=0.85 if text else 0.25,
        parser_provenance=provenance,
    )


def extract(filename: str, mime_type: str, data: bytes) -> tuple[str, dict[str, Any], float]:
    """Compatibility shim for the original stateless endpoint/tests."""
    result = sanitize_and_extract(filename, mime_type, data)
    return result.text, result.normalized_fields, result.confidence
